from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "software-total-experience" / "Software_TXD_HoloPass.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 103, 112)
BLACK = RGBColor(0, 0, 0)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "EAF6FA"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, bold=False, size=9.5, color=BLACK):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        for run in p.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_fill(hdr[i], HEADER_FILL)
        style_cell_text(hdr[i], bold=True, size=9.5, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            style_cell_text(cells[i], size=9)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        color, size = BLUE, 16
    elif level == 2:
        color, size = BLUE, 13
    else:
        color, size = DARK_BLUE, 12
    for run in p.runs:
        set_run_font(run, size=size, color=color, bold=True)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10.5, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=BLACK)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "HoloPass GS 2026 - Software & Total Experience Design"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Thiago Souza de Lima - RM 568732"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=8.5, color=MUTED)


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("HoloPass")
    set_run_font(r, size=26, color=BLACK, bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run("Software & Total Experience Design | Global Solution 2026 | Industria Espacial")
    set_run_font(r, size=13, color=MUTED)

    add_table(
        doc,
        ["Campo", "Informacao"],
        [
            ["Equipe", "Thiago Souza de Lima - RM 568732"],
            ["Curso", "Engenharia de Software - 1o Ano - FIAP"],
            ["Tema", "Pulseira inteligente com NFC, GNSS e planejamento de mobilidade"],
            ["Entrega", "Documento Word estruturado para Software & Total Experience Design"],
        ],
        [1900, 7460],
    )

    add_callout(
        doc,
        "Decisao de revisao",
        "A entrega revisada foca em produto viavel: diagnostico GNSS real pelo navegador, catraca NFC simulada, PWA offline, HoloRoute deterministico, proximo trem real quando a API cobre a estacao e cobertura urbana com OpenStreetMap/Overpass e GTFS SPTrans.",
    )

    add_heading(doc, "1. Sumario Executivo")
    add_body(
        doc,
        "O HoloPass e uma pulseira inteligente para transporte publico urbano. Ela une pagamento por NFC, localizacao por GNSS, planejamento de rota operacional e PWA offline para reduzir filas, inseguranca, dependencia do celular e falta de informacao durante deslocamentos.",
    )
    add_body(
        doc,
        "O projeto existe como sistema integrado: app web em HTML/CSS/JS puro, landing page, firmware Arduino, programa Python de menu, modelo matematico GNSS e documentacao de pitch. Cada disciplina entrega uma camada da mesma solucao.",
    )

    add_heading(doc, "2. Problema Real")
    add_body(
        doc,
        "Passageiros enfrentam catracas lentas, saldo pouco visivel, celulares expostos em estacoes lotadas e informacao limitada sobre rota, baldeacao e chegada. Pessoas com menor acesso a smartphone/dados moveis sofrem mais, ampliando desigualdade de mobilidade.",
    )
    add_body(
        doc,
        "O custo aparece em tempo perdido, risco de furto, atraso operacional e menor confianca no transporte publico. A solucao importa porque mobilidade previsivel melhora acesso a estudo, trabalho, saude e lazer.",
    )

    add_heading(doc, "3. Solucao Tecnologica")
    add_body(
        doc,
        "O HoloPass propoe uma pulseira com display, NFC, vibracao e leitura GNSS. O app demonstra autenticacao, saldo, recarga, historico, rotas, pagamento, modo offline, HoloRoute deterministico, proximo trem real quando disponivel e analise de areas mal atendidas.",
    )
    add_body(
        doc,
        "O HoloRoute monta a malha como grafo de estacoes, linhas e corredores de transferencia. No caso Osasco -> Trianon-MASP, a rota correta usa Linha 9, transferencia em Pinheiros para Linha 4, transferencia Paulista/Consolacao para Linha 2 e chegada em Trianon-MASP.",
    )
    add_body(
        doc,
        "No site principal, a antiga vitrine visual foi substituida por um painel operacional com diagnostico GNSS, validacao NFC simulada e metricas de viagem. Isso reforca funcionamento verificavel em vez de uma imagem estatica do produto.",
    )

    add_heading(doc, "4. Conexao com a Industria Espacial")
    add_body(
        doc,
        "A conexao espacial e direta por GNSS: o sistema usa latitude/longitude, precisao em metros e Haversine para identificar a estacao mais proxima. A camada de cobertura urbana usa OpenStreetMap/Overpass para localizar estacoes reais e GTFS SPTrans para contar paradas de onibus no raio consultado dentro da cobertura municipal de Sao Paulo. CBERS-4A/Amazonia-1 entram como contexto de Observacao da Terra, sem inventar imagem satelital real.",
    )

    add_heading(doc, "5. Arquitetura Integrada")
    add_table(
        doc,
        ["Camada", "Entrega", "Papel no sistema"],
        [
            ["Software/TXD", "Este documento", "Define problema, visao, arquitetura, backlog e fluxo"],
            ["Front-End Design", "Landing page", "Comunica problema, tecnologia, objetivos, publico e beneficios"],
            ["Web Development", "PWA principal", "Prototipo funcional de login, recarga, rota, NFC, GNSS, quiz e feedback"],
            ["Edge Computing", "Arduino", "Simula a pulseira fisica com GNSS, NFC, botoes, LEDs, buzzer, demanda urbana e telemetria"],
            ["Computational Thinking", "Python menu", "Demonstra logica de usuario, saldo, rota, historico e validacao"],
            ["DPS", "Modelo GNSS", "Explica matematica espacial com funcoes e graficos"],
            ["Storytelling", "Pitch", "Apresenta narrativa, proposta de valor, tecnologia e equipe"],
        ],
        [1800, 2100, 5460],
    )

    add_heading(doc, "6. Viabilidade Tecnica")
    add_bullet(doc, "A pulseira fisica pode ser prototipada com ESP32/Arduino, RFID/NFC, display OLED, motor de vibracao, LEDs de estado, botoes e bateria.")
    add_bullet(doc, "GNSS e NFC no Arduino sao simulados, mas a matematica de distancia e a decisao local sao reais.")
    add_bullet(doc, "O app web usa HTML, CSS e JavaScript puro, com PWA offline por service worker.")
    add_bullet(doc, "Proximo trem usa a API publica ViaMobilidade nas estacoes cobertas; fora disso, o app mostra fallback honesto.")
    add_bullet(doc, "A priorizacao urbana consulta Overpass/OpenStreetMap, usa GTFS SPTrans local e calcula distancia por Haversine.")
    add_bullet(doc, "Uma versao de producao exigiria dados oficiais, homologacao de pagamento e seguranca embarcada.")

    add_heading(doc, "7. Impacto e ODS")
    add_table(
        doc,
        ["ODS", "Conexao com o HoloPass", "Beneficio esperado"],
        [
            ["ODS 9", "Software, dispositivo fisico e dados de localizacao", "Infraestrutura de transporte mais inteligente"],
            ["ODS 10", "Menor dependencia de smartphone/dados moveis", "Acesso mais inclusivo a mobilidade"],
            ["ODS 11", "Rotas, baldeacoes e areas mal atendidas", "Cidades mais conectadas e previsiveis"],
            ["ODS 13", "Transporte coletivo mais atrativo", "Menor incentivo a deslocamento individual"],
        ],
        [1100, 3950, 4310],
    )

    add_heading(doc, "8. Declaracao da Visao do Produto")
    add_body(
        doc,
        "Para passageiros urbanos que precisam viajar com seguranca, previsibilidade e menos dependencia do celular, o HoloPass e uma pulseira de transporte inteligente que combina pagamento NFC, localizacao GNSS, rota operacional e alertas fisicos. Diferente de um bilhete comum, ele integra pagamento, posicao, historico e planejamento urbano em uma experiencia unica e acessivel.",
    )

    add_heading(doc, "9. Backlog Priorizado")
    add_table(
        doc,
        ["ID", "Historia de usuario", "Prioridade", "Criterio de aceite"],
        [
            ["US01", "Pagar por NFC na catraca sem tirar celular/cartao do bolso.", "Alta", "Pagamento debita saldo e registra historico."],
            ["US02", "Recarregar por PIX, cartao ou debito.", "Alta", "Valores R$20/50/100 atualizam saldo."],
            ["US03", "Detectar estacao por GNSS.", "Alta", "App mostra estacao mais proxima e precisao."],
            ["US04", "Calcular rota com baldeacoes reais.", "Alta", "Exibe linhas, paradas, tempo, tarifa e timeline."],
            ["US05", "Consultar saldo, historico e estatisticas.", "Media", "Painel atualiza apos recarga/pagamento."],
            ["US06", "Usar app offline como PWA.", "Media", "Service worker mantem app aberto sem rede."],
            ["US07", "Visualizar areas mal atendidas.", "Media", "GNSS/OSM mostram estacao proxima, distancia e prioridade."],
            ["US08", "Receber recomendacao HoloRoute explicavel.", "Media", "Painel mostra risco, lotacao e recomendacao deterministica."],
            ["US09", "Ver evidencias tecnicas da entrega.", "Alta", "Relatorio lista comando, resultado e pendencias."],
        ],
        [850, 3860, 1300, 3350],
    )

    add_heading(doc, "10. User Flow")
    add_table(
        doc,
        ["Etapa", "Acao do usuario", "Resposta do sistema"],
        [
            ["1", "Abre o HoloPass", "Restaura sessao ou mostra login/cadastro/demo."],
            ["2", "Consulta saldo", "Mostra saldo, historico e estatisticas."],
            ["3", "Recarrega se necessario", "Atualiza saldo por PIX/cartao/debito simulado."],
            ["4", "Detecta origem por GNSS ou seleciona manualmente", "Mostra estacao mais proxima e precisao."],
            ["5", "Escolhe destino", "HoloRoute calcula grafo de linhas e baldeacoes."],
            ["6", "Confere rota", "Mostra distancia operacional, tempo, tarifa, paradas e timeline."],
            ["7", "Aproxima NFC", "Debita saldo ou orienta recarga se insuficiente."],
            ["8", "Viaja", "Historico e PWA mantem dados essenciais disponiveis."],
        ],
        [900, 3600, 4860],
    )

    add_heading(doc, "11. Criterios de Aceite")
    for item in [
        "App sem erros de console nos fluxos principais.",
        "Rota Osasco -> Trianon-MASP usa Linha 9 -> Linha 4 -> Linha 2.",
        "Distancia exibida e rotulada como operacional estimada por trechos da rede.",
        "Nenhuma promessa tecnica sem evidencia ou API real nao comprovada.",
        "Edge compila e demonstra LED, buzzer, NFC, saldo e telemetria.",
        "CT Python usa conceitos obrigatorios e nao depende de pacote externo.",
        "Toda imagem local usada no app possui texto alternativo.",
        "Pendencias externas ficam separadas: video, Wokwi publico e organizacao GitHub.",
    ]:
        add_bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
